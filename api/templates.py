"""CV templates for HTML rendering and export."""
import io

TEMPLATES = {
    "modern": {
        "name": "Modern",
        "description": "Clean, minimalist design with accent colors",
        "css": """
            .cv-modern { font-family: 'Inter', sans-serif; max-width: 210mm; margin: 0 auto; padding: 40px; color: #1a1a2e; background: #fff; }
            .cv-modern .header { border-bottom: 3px solid #2563eb; padding-bottom: 16px; margin-bottom: 24px; }
            .cv-modern .name { font-size: 28px; font-weight: 800; color: #1a1a2e; margin: 0; }
            .cv-modern .contact { font-size: 13px; color: #64748b; margin-top: 6px; display: flex; gap: 16px; flex-wrap: wrap; }
            .cv-modern .section-title { font-size: 14px; font-weight: 700; color: #2563eb; text-transform: uppercase; letter-spacing: 1px; margin: 20px 0 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 4px; }
            .cv-modern .skill-tag { display: inline-block; background: #eff6ff; color: #2563eb; padding: 4px 12px; border-radius: 4px; font-size: 12px; margin: 3px; }
            .cv-modern .item { margin-bottom: 14px; }
            .cv-modern .item-title { font-weight: 700; font-size: 14px; color: #1a1a2e; }
            .cv-modern .item-subtitle { font-size: 12px; color: #64748b; }
            .cv-modern .item-desc { font-size: 13px; color: #334155; line-height: 1.5; margin-top: 4px; }
            .cv-modern ul { margin: 4px 0; padding-left: 18px; }
            .cv-modern li { font-size: 13px; color: #334155; margin-bottom: 3px; }
        """,
        "render": lambda data: _render_modern(data),
    },
    "classic": {
        "name": "Classic",
        "description": "Traditional resume layout, serif fonts",
        "css": """
            .cv-classic { font-family: 'Georgia', 'Times New Roman', serif; max-width: 210mm; margin: 0 auto; padding: 40px; color: #222; background: #fff; }
            .cv-classic .header { text-align: center; border-bottom: 2px solid #333; padding-bottom: 12px; margin-bottom: 20px; }
            .cv-classic .name { font-size: 26px; font-weight: 700; color: #000; margin: 0; letter-spacing: 1px; }
            .cv-classic .contact { font-size: 12px; color: #555; margin-top: 6px; }
            .cv-classic .section-title { font-size: 14px; font-weight: 700; color: #000; text-transform: uppercase; letter-spacing: 1.5px; margin: 18px 0 8px; border-bottom: 1px solid #999; padding-bottom: 4px; }
            .cv-classic .skill-tag { display: inline-block; background: #f0f0f0; color: #333; padding: 3px 10px; border-radius: 2px; font-size: 12px; margin: 2px; }
            .cv-classic .item { margin-bottom: 12px; }
            .cv-classic .item-title { font-weight: 700; font-size: 14px; color: #000; }
            .cv-classic .item-subtitle { font-size: 12px; color: #555; font-style: italic; }
            .cv-classic .item-desc { font-size: 13px; color: #333; line-height: 1.5; margin-top: 3px; }
            .cv-classic ul { margin: 3px 0; padding-left: 18px; }
            .cv-classic li { font-size: 13px; color: #333; margin-bottom: 2px; }
        """,
        "render": lambda data: _render_classic(data),
    },
    "minimal": {
        "name": "Minimal",
        "description": "No-frills, compact, ATS-friendly",
        "css": """
            .cv-minimal { font-family: 'Arial', 'Helvetica', sans-serif; max-width: 210mm; margin: 0 auto; padding: 30px; color: #111; background: #fff; font-size: 11px; }
            .cv-minimal .header { margin-bottom: 16px; }
            .cv-minimal .name { font-size: 20px; font-weight: 700; color: #000; margin: 0; }
            .cv-minimal .contact { font-size: 11px; color: #555; margin-top: 4px; }
            .cv-minimal .section-title { font-size: 11px; font-weight: 700; color: #000; margin: 12px 0 4px; border-bottom: 1px solid #ccc; padding-bottom: 2px; }
            .cv-minimal .skill-tag { display: inline; color: #111; font-size: 11px; margin: 0; }
            .cv-minimal .skill-tag::after { content: ", "; }
            .cv-minimal .skill-tag:last-child::after { content: ""; }
            .cv-minimal .item { margin-bottom: 8px; }
            .cv-minimal .item-title { font-weight: 700; font-size: 11px; color: #000; }
            .cv-minimal .item-subtitle { font-size: 11px; color: #555; }
            .cv-minimal .item-desc { font-size: 11px; color: #333; line-height: 1.4; margin-top: 2px; }
            .cv-minimal ul { margin: 2px 0; padding-left: 14px; }
            .cv-minimal li { font-size: 11px; color: #333; margin-bottom: 1px; }
        """,
        "render": lambda data: _render_minimal(data),
    },
}


def render_cv_html(cv_text, template="modern"):
    """Parse markdown-like CV text and render to HTML with selected template."""
    data = _parse_cv(cv_text)
    tmpl = TEMPLATES.get(template, TEMPLATES["modern"])
    body = tmpl["render"](data)
    css = tmpl["css"]
    
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<style>
{css}
@media print {{ body {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }} }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def _parse_cv(text):
    """Parse CV text into structured data."""
    lines = text.strip().split("\n")
    data = {
        "name": "Candidate",
        "contact": [],
        "summary": "",
        "skills": [],
        "experience": [],
        "education": [],
        "other": [],
    }
    
    current_section = "header"
    current_item = None
    contact_line = []
    
    import re
    
    for line in lines:
        clean = line.strip()
        if not clean:
            continue
        
        # Titles
        if clean.startswith("# ") and not clean.startswith("## "):
            data["name"] = clean.replace("# ", "").strip()
            current_section = "header"
            continue
        
        # Section headers
        if clean.startswith("## "):
            section = clean.replace("## ", "").strip().lower()
            if any(w in section for w in ["profile", "summary", "objective", "about"]):
                current_section = "summary"
            elif any(w in section for w in ["skill", "technology", "tool", "competenc"]):
                current_section = "skills"
            elif any(w in section for w in ["experience", "work", "employment", "professional"]):
                current_section = "experience"
                current_item = None
            elif any(w in section for w in ["education", "university", "college", "school", "degree"]):
                current_section = "education"
                current_item = None
            else:
                current_section = "other"
            continue
        
        # Sub-items (###)
        if clean.startswith("### "):
            sub = clean.replace("### ", "").strip()
            current_item = {"title": sub, "subtitle": "", "details": []}
            if current_section in data:
                data[current_section].append(current_item)
            continue
        
        # Bullet points
        if clean.startswith("- ") or clean.startswith("• ") or clean.startswith("* "):
            bullet = clean[2:].strip()
            if current_item and current_section in ["experience", "education"]:
                current_item["details"].append(bullet)
            elif current_section == "skills":
                data["skills"].append(bullet)
            else:
                data.setdefault("other", []).append(bullet)
            continue
        
        if clean.startswith("|"):
            contact_line = [c.strip() for c in clean.strip("|").split("|")]
            data["contact"] = [c for c in contact_line if c]
            continue
        
        # Plain text
        if current_section == "summary":
            if data["summary"]:
                data["summary"] += " " + clean
            else:
                data["summary"] = clean
        elif current_section == "skills":
            data["skills"].append(clean)
        elif current_item:
            current_item.setdefault("details", []).append(clean)
        else:
            data.setdefault("other", []).append(clean)
    
    return data


def _render_modern(data):
    parts = ["<div class='cv-modern'>"]
    parts.append("<div class='header'>")
    parts.append(f"<div class='name'>{data['name']}</div>")
    if data["contact"]:
        parts.append(f"<div class='contact'>{' | '.join(data['contact'][:5])}</div>")
    parts.append("</div>")
    
    if data["summary"]:
        parts.append(f"<div class='section-title'>Professional Profile</div>")
        parts.append(f"<p style='font-size:13px;color:#334155;line-height:1.5;'>{data['summary']}</p>")
    
    if data["skills"]:
        parts.append("<div class='section-title'>Skills</div>")
        parts.append("<div>" + "".join(f"<span class='skill-tag'>{s}</span>" for s in data["skills"]) + "</div>")
    
    if data["experience"]:
        parts.append("<div class='section-title'>Experience</div>")
        for exp in data["experience"]:
            parts.append("<div class='item'>")
            parts.append(f"<div class='item-title'>{exp['title']}</div>")
            if exp.get("subtitle"):
                parts.append(f"<div class='item-subtitle'>{exp['subtitle']}</div>")
            if exp.get("details"):
                parts.append("<ul>" + "".join(f"<li>{d}</li>" for d in exp["details"]) + "</ul>")
            parts.append("</div>")
    
    if data["education"]:
        parts.append("<div class='section-title'>Education</div>")
        for edu in data["education"]:
            parts.append("<div class='item'>")
            parts.append(f"<div class='item-title'>{edu['title']}</div>")
            if edu.get("subtitle"):
                parts.append(f"<div class='item-subtitle'>{edu['subtitle']}</div>")
            if edu.get("details"):
                parts.append("<ul>" + "".join(f"<li>{d}</li>" for d in edu["details"]) + "</ul>")
            parts.append("</div>")
    
    if data.get("other"):
        parts.append("<div class='section-title'>Additional</div>")
        parts.append("<ul>" + "".join(f"<li>{o}</li>" for o in data["other"]) + "</ul>")
    
    parts.append("</div>")
    return "\n".join(parts)


def _render_classic(data):
    parts = ["<div class='cv-classic'>"]
    parts.append("<div class='header'>")
    parts.append(f"<div class='name'>{data['name']}</div>")
    if data["contact"]:
        parts.append(f"<div class='contact'>{' | '.join(data['contact'][:5])}</div>")
    parts.append("</div>")
    
    if data["summary"]:
        parts.append(f"<div class='section-title'>Professional Profile</div>")
        parts.append(f"<p style='font-size:13px;color:#333;line-height:1.5;'>{data['summary']}</p>")
    
    if data["skills"]:
        parts.append("<div class='section-title'>Skills</div>")
        parts.append("<p>" + "".join(f"<span class='skill-tag'>{s}</span>" for s in data["skills"]) + "</p>")
    
    if data["experience"]:
        parts.append("<div class='section-title'>Experience</div>")
        for exp in data["experience"]:
            parts.append("<div class='item'>")
            parts.append(f"<div class='item-title'>{exp['title']}</div>")
            if exp.get("subtitle"):
                parts.append(f"<div class='item-subtitle'>{exp['subtitle']}</div>")
            if exp.get("details"):
                parts.append("<ul>" + "".join(f"<li>{d}</li>" for d in exp["details"]) + "</ul>")
            parts.append("</div>")
    
    if data["education"]:
        parts.append("<div class='section-title'>Education</div>")
        for edu in data["education"]:
            parts.append("<div class='item'>")
            parts.append(f"<div class='item-title'>{edu['title']}</div>")
            if edu.get("subtitle"):
                parts.append(f"<div class='item-subtitle'>{edu['subtitle']}</div>")
            if edu.get("details"):
                parts.append("<ul>" + "".join(f"<li>{d}</li>" for d in edu["details"]) + "</ul>")
            parts.append("</div>")
    
    if data.get("other"):
        parts.append("<div class='section-title'>Additional</div>")
        parts.append("<ul>" + "".join(f"<li>{o}</li>" for o in data["other"]) + "</ul>")
    
    parts.append("</div>")
    return "\n".join(parts)


def _render_minimal(data):
    parts = ["<div class='cv-minimal'>"]
    parts.append("<div class='header'>")
    parts.append(f"<div class='name'>{data['name']}</div>")
    if data["contact"]:
        parts.append(f"<div class='contact'>{' | '.join(data['contact'][:5])}</div>")
    parts.append("</div>")
    
    if data["summary"]:
        parts.append(f"<div class='section-title'>Profile</div>")
        parts.append(f"<div class='item-desc'>{data['summary']}</div>")
    
    if data["skills"]:
        parts.append("<div class='section-title'>Skills</div>")
        parts.append("<div>" + "".join(f"<span class='skill-tag'>{s}</span>" for s in data["skills"]) + "</div>")
    
    if data["experience"]:
        parts.append("<div class='section-title'>Experience</div>")
        for exp in data["experience"]:
            parts.append("<div class='item'>")
            parts.append(f"<div class='item-title'>{exp['title']}</div>")
            if exp.get("subtitle"):
                parts.append(f"<div class='item-subtitle'>{exp['subtitle']}</div>")
            if exp.get("details"):
                parts.append("<ul>" + "".join(f"<li>{d}</li>" for d in exp["details"]) + "</ul>")
            parts.append("</div>")
    
    if data["education"]:
        parts.append("<div class='section-title'>Education</div>")
        for edu in data["education"]:
            parts.append("<div class='item'>")
            parts.append(f"<div class='item-title'>{edu['title']}</div>")
            if edu.get("subtitle"):
                parts.append(f"<div class='item-subtitle'>{edu['subtitle']}</div>")
            if edu.get("details"):
                parts.append("<ul>" + "".join(f"<li>{d}</li>" for d in edu["details"]) + "</ul>")
            parts.append("</div>")
    
    if data.get("other"):
        parts.append("<div class='section-title'>Additional</div>")
        parts.append("<div>" + " | ".join(data["other"][:8]) + "</div>")
    
    parts.append("</div>")
    return "\n".join(parts)


def get_template_list():
    return [
        {"id": k, "name": v["name"], "description": v["description"]}
        for k, v in TEMPLATES.items()
    ]


def export_to_html(cv_text, template="modern"):
    return render_cv_html(cv_text, template)


def export_to_docx(cv_text, template="modern"):
    """Export CV to python-docx format."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    
    data = _parse_cv(cv_text)
    
    # Name
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(data["name"])
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    name_para.space_after = Pt(4)
    
    # Contact
    if data["contact"]:
        contact_para = doc.add_paragraph()
        contact_run = contact_para.add_run(" | ".join(data["contact"][:5]))
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        contact_para.space_after = Pt(12)
    
    # Separator line
    doc.add_paragraph("_" * 60)
    
    # Summary
    if data["summary"]:
        _add_docx_section_title(doc, "Professional Profile")
        p = doc.add_paragraph(data["summary"])
        p.space_after = Pt(6)
    
    # Skills
    if data["skills"]:
        _add_docx_section_title(doc, "Skills")
        skills_text = ", ".join(data["skills"])
        p = doc.add_paragraph(skills_text)
        p.space_after = Pt(6)
    
    # Experience
    if data["experience"]:
        _add_docx_section_title(doc, "Experience")
        for exp in data["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(exp["title"])
            run.bold = True
            run.font.size = Pt(12)
            p.space_after = Pt(2)
            if exp.get("details"):
                for detail in exp["details"]:
                    dp = doc.add_paragraph(detail, style="List Bullet")
                    dp.space_after = Pt(1)
    
    # Education
    if data["education"]:
        _add_docx_section_title(doc, "Education")
        for edu in data["education"]:
            p = doc.add_paragraph()
            run = p.add_run(edu["title"])
            run.bold = True
            run.font.size = Pt(12)
            p.space_after = Pt(2)
            if edu.get("details"):
                for detail in edu["details"]:
                    dp = doc.add_paragraph(detail, style="List Bullet")
                    dp.space_after = Pt(1)
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_docx_section_title(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    p.space_before = Pt(12)
    p.space_after = Pt(4)

